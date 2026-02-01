"""
Document Parser Module.

Provides unified parsing for multiple document formats:
- PDF (with OCR fallback)
- DOCX (Microsoft Word)
- HTML/Web pages
- Plain text
- Markdown

Extracts text content, metadata, and structure information.
"""

import io
import logging
import mimetypes
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, BinaryIO

import chardet

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    HTML = "html"
    MARKDOWN = "markdown"
    TEXT = "text"
    UNKNOWN = "unknown"


@dataclass
class ParsedSection:
    """A section/element from a parsed document."""
    content: str
    section_type: str  # heading, paragraph, table, list, code, etc.
    level: int = 0  # For headings (h1=1, h2=2, etc.)
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Result of document parsing."""
    doc_id: str
    filename: str
    doc_type: DocumentType
    content: str  # Full text content
    sections: list[ParsedSection] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    page_count: int = 0
    word_count: int = 0
    char_count: int = 0
    parsed_at: datetime = field(default_factory=datetime.utcnow)
    
    def __post_init__(self):
        """Calculate statistics after init."""
        if self.content:
            self.char_count = len(self.content)
            self.word_count = len(self.content.split())


class PDFParser:
    """
    PDF document parser using PyMuPDF (fitz).
    
    Features:
    - Text extraction with layout preservation
    - Table detection
    - Image extraction (optional)
    - OCR fallback for scanned documents
    """
    
    def __init__(self, use_ocr_fallback: bool = True):
        self.use_ocr_fallback = use_ocr_fallback
    
    def parse(
        self,
        file_path: str | Path | None = None,
        file_content: bytes | BinaryIO | None = None,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """
        Parse a PDF document.
        
        Args:
            file_path: Path to the PDF file
            file_content: PDF file content as bytes or file-like object
            doc_id: Document identifier
        
        Returns:
            ParsedDocument with extracted content
        """
        import fitz  # PyMuPDF
        
        # Open document
        if file_path:
            doc = fitz.open(file_path)
            filename = Path(file_path).name
        elif file_content:
            if isinstance(file_content, bytes):
                doc = fitz.open(stream=file_content, filetype="pdf")
            else:
                doc = fitz.open(stream=file_content.read(), filetype="pdf")
            filename = "uploaded.pdf"
        else:
            raise ValueError("Either file_path or file_content must be provided")
        
        doc_id = doc_id or filename
        
        # Extract content
        full_text = []
        sections = []
        
        for page_num, page in enumerate(doc, start=1):
            # Extract text with layout
            text = page.get_text("text")
            full_text.append(text)
            
            # Try to identify structure
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            content = span.get("text", "").strip()
                            if content:
                                # Detect section type based on font size
                                font_size = span.get("size", 12)
                                if font_size >= 18:
                                    section_type = "heading"
                                    level = 1
                                elif font_size >= 14:
                                    section_type = "heading"
                                    level = 2
                                else:
                                    section_type = "paragraph"
                                    level = 0
                                
                                sections.append(ParsedSection(
                                    content=content,
                                    section_type=section_type,
                                    level=level,
                                    metadata={"page": page_num, "font_size": font_size},
                                ))
        
        # Extract metadata
        metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "creation_date": doc.metadata.get("creationDate", ""),
            "modification_date": doc.metadata.get("modDate", ""),
        }
        
        content = "\n".join(full_text)
        
        # OCR fallback if text is too short
        if self.use_ocr_fallback and len(content.strip()) < 100:
            logger.info(f"Text too short, attempting OCR for {filename}")
            content = self._ocr_fallback(doc)
        
        doc.close()
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            doc_type=DocumentType.PDF,
            content=content,
            sections=sections,
            metadata=metadata,
            page_count=len(doc),
        )
    
    def _ocr_fallback(self, doc) -> str:
        """Use OCR for scanned PDFs."""
        try:
            import pytesseract
            from PIL import Image
            
            full_text = []
            for page in doc:
                # Render page to image
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # OCR
                text = pytesseract.image_to_string(img)
                full_text.append(text)
            
            return "\n".join(full_text)
        except ImportError:
            logger.warning("pytesseract not installed, OCR unavailable")
            return ""


class DOCXParser:
    """
    Microsoft Word document parser using python-docx.
    
    Features:
    - Text extraction with structure
    - Table extraction
    - Header/footer handling
    - Metadata extraction
    """
    
    def parse(
        self,
        file_path: str | Path | None = None,
        file_content: bytes | BinaryIO | None = None,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """Parse a DOCX document."""
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # Open document
        if file_path:
            doc = Document(file_path)
            filename = Path(file_path).name
        elif file_content:
            if isinstance(file_content, bytes):
                doc = Document(io.BytesIO(file_content))
            else:
                doc = Document(file_content)
            filename = "uploaded.docx"
        else:
            raise ValueError("Either file_path or file_content must be provided")
        
        doc_id = doc_id or filename
        
        # Extract content
        full_text = []
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            full_text.append(text)
            
            # Determine section type
            style_name = para.style.name.lower() if para.style else ""
            
            if "heading 1" in style_name:
                section_type = "heading"
                level = 1
            elif "heading 2" in style_name:
                section_type = "heading"
                level = 2
            elif "heading 3" in style_name:
                section_type = "heading"
                level = 3
            elif "title" in style_name:
                section_type = "heading"
                level = 1
            elif "list" in style_name:
                section_type = "list"
                level = 0
            else:
                section_type = "paragraph"
                level = 0
            
            sections.append(ParsedSection(
                content=text,
                section_type=section_type,
                level=level,
                metadata={"style": para.style.name if para.style else ""},
            ))
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            
            table_content = "\n".join(table_text)
            full_text.append(table_content)
            
            sections.append(ParsedSection(
                content=table_content,
                section_type="table",
                level=0,
                metadata={"table_index": table_idx},
            ))
        
        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            doc_type=DocumentType.DOCX,
            content="\n\n".join(full_text),
            sections=sections,
            metadata=metadata,
        )


class HTMLParser:
    """
    HTML document parser using BeautifulSoup.
    
    Features:
    - Clean text extraction
    - Structure preservation
    - Script/style removal
    - Link extraction
    """
    
    def __init__(self, parser: str = "html.parser"):
        self.parser = parser
    
    def parse(
        self,
        file_path: str | Path | None = None,
        file_content: str | bytes | None = None,
        url: str | None = None,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """Parse an HTML document."""
        from bs4 import BeautifulSoup
        
        # Get HTML content
        if file_path:
            with open(file_path, "rb") as f:
                content = f.read()
            filename = Path(file_path).name
        elif file_content:
            content = file_content
            filename = "uploaded.html"
        elif url:
            import requests
            response = requests.get(url, timeout=30)
            content = response.content
            filename = url
        else:
            raise ValueError("Either file_path, file_content, or url must be provided")
        
        # Decode if bytes
        if isinstance(content, bytes):
            detected = chardet.detect(content)
            encoding = detected.get("encoding", "utf-8")
            content = content.decode(encoding, errors="ignore")
        
        doc_id = doc_id or filename
        
        # Parse HTML
        soup = BeautifulSoup(content, self.parser)
        
        # Remove unwanted elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
        
        # Extract sections
        sections = []
        
        # Headings
        for level in range(1, 7):
            for heading in soup.find_all(f"h{level}"):
                text = heading.get_text(strip=True)
                if text:
                    sections.append(ParsedSection(
                        content=text,
                        section_type="heading",
                        level=level,
                    ))
        
        # Paragraphs
        for para in soup.find_all("p"):
            text = para.get_text(strip=True)
            if text:
                sections.append(ParsedSection(
                    content=text,
                    section_type="paragraph",
                    level=0,
                ))
        
        # Lists
        for ul in soup.find_all(["ul", "ol"]):
            items = [li.get_text(strip=True) for li in ul.find_all("li")]
            if items:
                sections.append(ParsedSection(
                    content="\n".join(f"• {item}" for item in items),
                    section_type="list",
                    level=0,
                ))
        
        # Code blocks
        for code in soup.find_all(["pre", "code"]):
            text = code.get_text(strip=True)
            if text:
                sections.append(ParsedSection(
                    content=text,
                    section_type="code",
                    level=0,
                ))
        
        # Full text
        full_text = soup.get_text(separator="\n", strip=True)
        
        # Clean up excessive whitespace
        full_text = re.sub(r"\n{3,}", "\n\n", full_text)
        
        # Extract metadata
        metadata = {
            "title": soup.title.string if soup.title else "",
            "url": url or "",
        }
        
        # Meta tags
        for meta in soup.find_all("meta"):
            name = meta.get("name", meta.get("property", ""))
            content = meta.get("content", "")
            if name and content:
                metadata[name] = content
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            doc_type=DocumentType.HTML,
            content=full_text,
            sections=sections,
            metadata=metadata,
        )


class MarkdownParser:
    """
    Markdown document parser.
    
    Features:
    - Heading extraction
    - Code block handling
    - Link extraction
    """
    
    def parse(
        self,
        file_path: str | Path | None = None,
        file_content: str | None = None,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """Parse a Markdown document."""
        # Get content
        if file_path:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            filename = Path(file_path).name
        elif file_content:
            content = file_content
            filename = "uploaded.md"
        else:
            raise ValueError("Either file_path or file_content must be provided")
        
        doc_id = doc_id or filename
        
        # Extract sections
        sections = []
        
        # Headings
        heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
        for match in heading_pattern.finditer(content):
            level = len(match.group(1))
            text = match.group(2).strip()
            sections.append(ParsedSection(
                content=text,
                section_type="heading",
                level=level,
            ))
        
        # Code blocks
        code_pattern = re.compile(r"```(\w*)\n(.*?)```", re.DOTALL)
        for match in code_pattern.finditer(content):
            language = match.group(1) or "text"
            code = match.group(2).strip()
            sections.append(ParsedSection(
                content=code,
                section_type="code",
                level=0,
                metadata={"language": language},
            ))
        
        # Extract title from first heading
        title = ""
        title_match = re.match(r"^#\s+(.+)$", content, re.MULTILINE)
        if title_match:
            title = title_match.group(1).strip()
        
        metadata = {
            "title": title,
        }
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            doc_type=DocumentType.MARKDOWN,
            content=content,
            sections=sections,
            metadata=metadata,
        )


class TextParser:
    """Plain text document parser."""
    
    def parse(
        self,
        file_path: str | Path | None = None,
        file_content: str | bytes | None = None,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """Parse a plain text document."""
        if file_path:
            with open(file_path, "rb") as f:
                content = f.read()
            filename = Path(file_path).name
        elif file_content:
            content = file_content
            filename = "uploaded.txt"
        else:
            raise ValueError("Either file_path or file_content must be provided")
        
        # Decode if bytes
        if isinstance(content, bytes):
            detected = chardet.detect(content)
            encoding = detected.get("encoding", "utf-8")
            content = content.decode(encoding, errors="ignore")
        
        doc_id = doc_id or filename
        
        # Split into paragraphs
        paragraphs = re.split(r"\n{2,}", content.strip())
        sections = [
            ParsedSection(content=p.strip(), section_type="paragraph", level=0)
            for p in paragraphs if p.strip()
        ]
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            doc_type=DocumentType.TEXT,
            content=content,
            sections=sections,
            metadata={},
        )


class UnifiedDocumentParser:
    """
    Unified document parser that auto-detects file type.
    
    Example:
        parser = UnifiedDocumentParser()
        doc = parser.parse("contract.pdf")
        doc = parser.parse_content(file_bytes, filename="doc.docx")
    """
    
    def __init__(self):
        self._pdf_parser = PDFParser()
        self._docx_parser = DOCXParser()
        self._html_parser = HTMLParser()
        self._markdown_parser = MarkdownParser()
        self._text_parser = TextParser()
        
        # Extension to parser mapping
        self._parsers = {
            ".pdf": self._pdf_parser,
            ".docx": self._docx_parser,
            ".doc": self._docx_parser,
            ".html": self._html_parser,
            ".htm": self._html_parser,
            ".md": self._markdown_parser,
            ".markdown": self._markdown_parser,
            ".txt": self._text_parser,
        }
    
    def detect_type(self, filename: str) -> DocumentType:
        """Detect document type from filename."""
        ext = Path(filename).suffix.lower()
        
        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOC,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".txt": DocumentType.TEXT,
        }
        
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    def parse(
        self,
        file_path: str | Path,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """
        Parse a document from file path.
        
        Args:
            file_path: Path to the document
            doc_id: Optional document identifier
        
        Returns:
            ParsedDocument with extracted content
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        parser = self._parsers.get(ext)
        if not parser:
            logger.warning(f"Unknown file type: {ext}, using text parser")
            parser = self._text_parser
        
        return parser.parse(file_path=file_path, doc_id=doc_id)
    
    def parse_content(
        self,
        content: bytes | str,
        filename: str,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """
        Parse document from content bytes.
        
        Args:
            content: Document content
            filename: Original filename (for type detection)
            doc_id: Optional document identifier
        
        Returns:
            ParsedDocument with extracted content
        """
        ext = Path(filename).suffix.lower()
        
        parser = self._parsers.get(ext)
        if not parser:
            logger.warning(f"Unknown file type: {ext}, using text parser")
            parser = self._text_parser
        
        return parser.parse(file_content=content, doc_id=doc_id)
    
    def parse_url(
        self,
        url: str,
        doc_id: str | None = None,
    ) -> ParsedDocument:
        """Parse a web page from URL."""
        return self._html_parser.parse(url=url, doc_id=doc_id or url)


# Singleton instance
_parser: UnifiedDocumentParser | None = None


def get_document_parser() -> UnifiedDocumentParser:
    """Get the singleton document parser."""
    global _parser
    if _parser is None:
        _parser = UnifiedDocumentParser()
    return _parser


__all__ = [
    "DocumentType",
    "ParsedDocument",
    "ParsedSection",
    "PDFParser",
    "DOCXParser",
    "HTMLParser",
    "MarkdownParser",
    "TextParser",
    "UnifiedDocumentParser",
    "get_document_parser",
]
